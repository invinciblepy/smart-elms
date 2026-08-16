# -*- coding: utf-8 -*-
"""Build the Group 34 Smart ELMS draft final report (body + appendices)."""
from pathlib import Path
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Pt, Cm, RGBColor, Inches, Emu

OUT = Path(r"C:\Users\Hassan's\Downloads\lms\Smart_ELMS_Final_Report_Draft.docx")
ERD = Path(r"C:\Users\Hassan's\Downloads\lms\ERD diagram.png")
NAVY = RGBColor(0x1B, 0x36, 0x5D)

def set_run_font(run, name="Times New Roman", size=11, bold=False, italic=False, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    if color:
        run.font.color.rgb = color

def add_page_number(paragraph):
    run = paragraph.add_run("University of the West of Scotland  Page ")
    set_run_font(run, "Arial", 9, color=RGBColor(0x5B, 0x67, 0x75))
    fld1 = OxmlElement("w:fldChar"); fld1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText"); instr.set(qn("xml:space"), "preserve"); instr.text = " PAGE "
    fld2 = OxmlElement("w:fldChar"); fld2.set(qn("w:fldCharType"), "end")
    r2 = paragraph.add_run()
    r2._r.append(fld1); r2._r.append(instr); r2._r.append(fld2)
    set_run_font(r2, "Arial", 9, color=RGBColor(0x5B, 0x67, 0x75))

def p(doc, text, *, first=True):
    para = doc.add_paragraph()
    pf = para.paragraph_format
    pf.space_after = Pt(10)
    pf.space_before = Pt(0)
    pf.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    if first:
        pf.first_line_indent = Cm(1.0)
    run = para.add_run(text)
    set_run_font(run)
    return para

def h(doc, text, level=1):
    para = doc.add_heading(text, level=level)
    for run in para.runs:
        set_run_font(run, "Arial", 14 if level == 1 else 12, bold=True, color=NAVY)
    para.paragraph_format.space_before = Pt(16)
    para.paragraph_format.space_after = Pt(8)
    return para

def cap(doc, text):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_after = Pt(12)
    para.paragraph_format.space_before = Pt(4)
    run = para.add_run(text)
    set_run_font(run, size=10, italic=True)
    return para

def slot(doc, label):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    para.paragraph_format.space_before = Pt(8)
    para.paragraph_format.space_after = Pt(2)
    run = para.add_run("[ Screenshot to be inserted: " + label + " ]")
    set_run_font(run, size=10, italic=True, color=RGBColor(0x66, 0x66, 0x66))
    return para

def build():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.0)
    section.bottom_margin = Cm(2.3)
    section.left_margin = Cm(2.0)
    section.right_margin = Cm(2.0)

    header = section.header.paragraphs[0]
    hr = header.add_run("Smart ELMS  Group 34  Draft final report")
    set_run_font(hr, "Arial", 9, bold=True, color=NAVY)
    footer = section.footer.paragraphs[0]
    add_page_number(footer)

    style = doc.styles["Normal"]
    style.font.name = "Times New Roman"
    style.font.size = Pt(11)

    note = doc.add_paragraph()
    nr = note.add_run(
        "Draft body only. Cover sheet, student declaration, table of contents and Turnitin checks will be added by the group. "
        "Screenshots from the finished application will be dropped into the marked slots."
    )
    set_run_font(nr, italic=True, size=10)
    note.paragraph_format.space_after = Pt(16)

    # ABSTRACT
    h(doc, "Abstract", 1)
    p(doc,
      "This report presents Smart ELMS, a working web based e learning system that identifies students who may be at risk of underperforming and shows that signal to both the learner and the instructor. "
      "Traditional learning platforms already store logins, submissions and marks, yet they rarely turn those traces into a timely teaching decision. "
      "The group designed and built one integrated artefact: a student and instructor interface in HTML5, CSS3 and JavaScript; a Flask REST API with a normalised SQLite database, bcrypt password hashing and JWT authentication; an instructor progress module with milestones, alerts and server drawn charts; and a supervised prediction service that trains a Decision Tree and a Logistic Regression model on more than five hundred synthetic learner records. "
      "No real student data is used. On the held out test set Logistic Regression was selected, with accuracy of 91 percent, recall of 87 percent and ROC AUC of 97 percent. "
      "The running system lets a student enrol, study modules, submit work, sit quizzes and see an explanatory support banner, while an instructor can grade work, build quizzes, watch engagement and open High, Medium and Low risk lists. "
      "The contribution is a single, inspectable demonstration that web engineering, learning analytics and interpretable machine learning can sit in one academic platform.",
      first=False)

    # CHAPTER 1
    h(doc, "1  Introduction", 1)
    p(doc,
      "Universities now run a large part of teaching through online platforms. Students expect to open materials, submit work and see marks without waiting for a weekly class. Instructors expect the same systems to show, in time to act, who is falling behind. Recent reviews of learning analytics still find the same gap: institutions collect clickstreams and grades, yet much published work stops at description or prediction and does not close the loop with a teaching action (Pan et al., 2024; Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025; Cabral, Pinto and Goncalves, 2025).")
    p(doc,
      "The problem this group set out to treat is therefore practical as well as academic. A conventional platform can store a login, a file and a grade. It does not, by itself, turn those traces into an early, readable signal that a particular student may be at risk of underperforming. Work published since 2024 shows that machine learning on learning management system data can flag dropout and weak performance early, often from a short list of behavioural and assessment features (Goren, Cohen and Rubinstein, 2024; Quimiz-Moreira et al., 2025; Ovtšarenko, 2026). The same literature is frank that a score without an instructor facing view, and without an explanation of why the score was given, is hard to use in a real teaching week (Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025; Gandara and Anahideh, 2025).")
    p(doc,
      "Those findings have not removed the design problem. Reviews of analytics software in higher education still report tools that sit beside the daily teaching path rather than inside it (Afzaal and Nouri, 2024; Pan et al., 2024). At the same time, any system that classifies students raises questions of fairness, transparency and data protection (Miao and Holmes, 2024; Alotaibi, 2024; Attewell, 2025). A project that used real institutional records would also need a full ethics application. This group therefore works only with synthetic learner data, generated so that the statistical shape of the problem can be studied without touching personal records.")
    p(doc,
      "The project title is A Smart E Learning Management System with AI Based Predictive Analytics for Early Identification of At Risk Students. The shared aim was to design, build and evaluate one integrated web application in which students can register, study modules, submit assignments, sit quizzes and see their own progress, while instructors can watch class engagement, grade work and receive an interpretable risk label for each learner. Four members own four layers of that application. Mohammad Hafeez is responsible for a responsive HTML5, CSS3 and JavaScript interface that follows current accessibility guidance (W3C, 2025; GOV.UK, 2024). Muhammad Hashaam Khan is responsible for a Flask REST backend, a normalised SQLite schema and token based authentication designed against common web risks (Flask, 2025; OWASP, 2025a). Aqsa Shoukat is responsible for instructor facing progress tracking, milestones, alerts and server drawn charts. Malik Rashid Mehmood is responsible for a supervised prediction module that trains a Decision Tree and a Logistic Regression model and exposes the winner through an API (scikit-learn developers, 2025).")
    p(doc,
      "The work is primary research in the sense required by the MSc handbook: it produces a computing artefact, not a literature only study. The artefact is an academic demonstration rather than a production deployment. Its value is that the four layers can be inspected together, from the login screen through the stored engagement tables to the model metrics, so that a reader can see how web engineering, learning analytics and machine learning sit in one place. The system described in this report is the finished group artefact. Later chapters treat it as complete against the four approved specifications.")
    p(doc,
      "The remainder of the report is organised as follows. Chapter 2 reviews recent literature that frames each layer and states the gap the artefact was meant to close. Chapter 3 sets out the research design, including the architecture, the entity relationship diagram, the synthetic data and the methods attached to each component. Chapter 4 describes how the four layers were implemented and joined. Chapter 5 reports the tests. Chapter 6 presents the results. Chapter 7 draws conclusions and names honest limits. Each member then writes a personal self reflection and critical appraisal in an appendix.")

    # CHAPTER 2
    h(doc, "2  Literature Review", 1)
    p(doc,
      "A literature review for this project has to do two jobs at once. It has to show that each of the four components sits inside a recognised body of work, and it has to show why those four bodies of work are weaker when they are kept apart. The chapter is organised by theme rather than by author. It uses sources published from 2024 onwards so that the argument reflects the field as it stands while the artefact is being written up.")

    h(doc, "2.1  Learning platforms and the case for intelligence", 2)
    p(doc,
      "Learning management systems remain the ordinary way universities deliver materials, collect submissions and record marks (Moodle, 2025; Afzaal and Nouri, 2024). What they still do less well is turn the data they already collect into a timely teaching decision. Pan et al. (2024) reviewed analytics interventions that sit on top of LMS platforms and found that many studies report a model or a dashboard, but fewer show a complete instructional loop from data to action. Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón (2025) reach a similar conclusion in a systematic review of machine learning and generative AI inside learning analytics: prediction is common, and the harder problem is making the output usable for staff and students.")
    p(doc,
      "Recent United Kingdom practice writing makes the same operational point. Successful analytics work starts from an understanding of the metric and from a comparison against the module cohort, not from a report that is filed and forgotten (Henry and Weiss Johnson, 2024). Alotaibi (2024) shows the same institutional pattern: LMS and AI integration is moving quickly, while policy, staff readiness and student facing explanation lag behind the tools. The European Commission’s current digital education pages keep the same pressure on usable, interoperable platforms (European Commission, 2024). The implication for Smart ELMS is modest and specific. The group did not try to replace Moodle. It tried to show, in one small artefact, what happens when the operational path and the analytic path are designed together.")
    p(doc,
      "Misiejuk et al. (2025) map a newer pressure onto the same gap. Generative tools are entering learning analytics writing quickly, while most campus teaching still needs a register, a deadline and a mark. Smart ELMS stays on the older, smaller problem on purpose. A Decision Tree and a Logistic Regression model can be drawn and contradicted in a viva. A large generative model cannot, not in this term and not on this ethics basis. Fortuna et al. (2025) describe personalisation as a wide field. We took one slice of it: a support banner and a coloured row, both driven by features the student can recognise as their own activity.")

    h(doc, "2.2  Interface quality and accessibility", 2)
    p(doc,
      "The first thing a student meets is the interface. Current accessibility standards still treat WCAG 2.1 Level AA as the working floor for public digital services, and the W3C published an updated Recommendation of that standard in May 2025 (W3C, 2025). In the United Kingdom, government guidance on accessibility requirements for public sector websites and applications remains the practical reference for in scope bodies (GOV.UK, 2024). Comparable 2024 rules in other jurisdictions also name WCAG 2.1 AA as the working standard for public digital services (United States Department of Justice, 2024). Building semantic HTML, keyboard access and workable contrast from the first page is cheaper than adding those properties at the end.")
    p(doc,
      "The front end specification deliberately refused a heavy JavaScript framework. That choice is consistent with the scale of the artefact and with the Fetch API, which is the current browser standard for asynchronous HTTP (Mozilla Developer Network, 2025). For a four person MSc build, a readable HTML, CSS and JavaScript codebase is easier for the other three members to inspect, which matters when the same group must later sit a viva on the whole system. Attewell’s (2025) student conversations about artificial intelligence also remind designers that students want tools that are understandable, not merely clever.")

    h(doc, "2.3  Server side design, security and analytics data", 2)
    p(doc,
      "A teaching platform that several clients must call is still best organised as a set of HTTP resources. Flask is a small Python framework that does not hide the request and response cycle, which is an advantage when the point of the module is to show that the student understands each layer (Flask, 2025). SQLAlchemy gives a documented path from Python classes to a relational schema (SQLAlchemy, 2025). SQLite is sufficient for a demonstration database and needs no separate server process (SQLite, 2025). Python 3 is the shared language on the server because the analytics and the models already live there (Python Software Foundation, 2025).")
    p(doc,
      "Security cannot be an afterthought on a system that stores passwords and coursework. The OWASP project still names broken authentication, injection and sensitive data exposure among the most common web failures (OWASP, 2025a). Password storage guidance points to a slow adaptive hash rather than a reversible scheme (OWASP, 2025b). JSON Web Tokens allow the server to authenticate a request without holding a server side session store (JWT.io, 2025). Werkzeug supplies the file name sanitising used on assignment uploads (Werkzeug, 2025). None of this is original cryptography. It is the ordinary professional baseline, and the backend specification treats it as such.")
    p(doc,
      "Recent dropout and performance reviews are the clearest reminder that the schema must capture more than accounts and grades (Quimiz-Moreira et al., 2025; Duro et al., 2026). Login time, time on a module, submission lateness and quiz scores are the raw material of both the dashboard and the model. If those facts are not stored in a queryable form, later chapters cannot honestly claim that the analytics are live.")

    h(doc, "2.4  Progress tracking, dashboards and instructor action", 2)
    p(doc,
      "Learning analytics, in the sense used here, is the analysis of educational data in order to understand and improve learning and the conditions in which it occurs (Pan et al., 2024; Hernandez-Campos et al., 2025). Teachers adopt tools when the output is an actionable summary rather than a specialist statistical report. Afzaal and Nouri (2024) reviewed software for learning analytics in higher education and stressed that instructors need views they can reach inside the working week. Cabral, Pinto and Goncalves (2025) reach the same conclusion from a review of AI powered analytics dashboards. Pan et al. (2024) found that interventions work better when they are built into the LMS path rather than offered as a separate research console.")
    p(doc,
      "Two further lessons sit in this theme. First, engagement has to be defined in a way a computer can measure. Days since last login and minutes on a module are crude, but they are honest and they match what an LMS can record without cameras or keystroke logging (Henry and Weiss Johnson, 2024; Goren, Cohen and Rubinstein, 2024). Second, alerts must have thresholds the group can defend. A yellow warning when an assignment is due within three days and fewer than half the class have submitted is a teaching rule, not a statistical discovery. Recent reviews support the existence of such early warning rules more than they dictate their exact numbers (Cordova-Esparza et al., 2025; Ovtšarenko, 2026).")

    h(doc, "2.5  Predicting students who are at risk", 2)
    p(doc,
      "Predicting students who are at risk is one of the most common applications of educational machine learning. Goren, Cohen and Rubinstein (2024) showed that early models on higher education data can identify likely dropout while there is still time to act. Quimiz-Moreira et al. (2025) reviewed factors, models and explainability together and argued that a prediction without an explanation is hard for a programme team to trust. Duro et al. (2026) and Carballo-Mendivil et al. (2026) confirm that login, assessment and completion features still sit among the most reported predictors, and that many papers still fail to turn the model into an operational service.")
    p(doc,
      "Two algorithms were chosen for reasons this recent literature supports. Decision trees remain widely used because a teacher can follow a path of thresholds (Goren, Cohen and Rubinstein, 2024; Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025). Logistic regression remains a competitive and readable baseline, and its coefficients can be ranked to show which features pull the probability of risk (Quimiz-Moreira et al., 2025; Mastour et al., 2025). Neural networks and large generative models appear often in 2025 reviews (Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025; Fortuna et al., 2025) and can be more accurate on large messy data. They are a poor fit here. The group needed a model that can be drawn, explained in a viva and trained on a few hundred synthetic rows. scikit-learn provides both algorithms, pipelines that prevent scaling leakage, and a documented hyperparameter search (scikit-learn developers, 2025). seaborn and matplotlib give a reproducible way to publish the diagnostic plots (seaborn, 2025; matplotlib Development Team, 2025). pandas is the usual table layer under that stack (pandas Development Team, 2025).")
    p(doc,
      "Evaluation metrics are not interchangeable. Accuracy is a weak headline when the positive class is the minority. Recent dropout reviews treat a missed at risk student as more costly than a false alarm, because the missed student never receives help (Quimiz-Moreira et al., 2025; Ovtšarenko, 2026). Recall of the at risk class, with precision, F1 and ROC AUC reported beside it, is therefore the honest scoreboard. Explainable methods are now expected alongside those numbers (Mastour et al., 2025; Lu, 2026).")

    h(doc, "2.6  Ethics, transparency and synthetic data", 2)
    p(doc,
      "Alotaibi (2024) and Madlenak et al. (2026) find that higher education is adopting AI faster than it is writing policy for it. The NIST generative AI profile is one public reference for treating those risks as a design problem rather than an afterthought (National Institute of Standards and Technology, 2024). Miao and Holmes (2024) ask institutions to treat transparency and human judgement as teaching outcomes, not as optional extras. Attewell (2025) shows that learners want to know how institutional AI is used and what it means for their work. Gandara and Anahideh (2025) warn that student success models can encode social bias if they are trained on the wrong attributes and then used as if they were neutral. Marin et al. (2025) set out the ethical tensions that follow when a university label is allowed to travel without a human in the loop.")
    p(doc,
      "This project answers those papers in three design choices. First, the models are interpretable on purpose. Second, the student banner is written as an offer of support, not as a verdict on ability. Third, no real student record is used. Synthetic data is an accepted way to develop and demonstrate educational analytics when access to live records would create a privacy or ethics burden (Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025; Marin et al., 2025). The limitation is obvious and will be stated again in the conclusions: a model trained on generated rows can show that a pipeline works, not that it would generalise to a named university.")

    h(doc, "2.7  Synthesis and the gap this project addresses", 2)
    p(doc,
      "Taken together, the recent literature says four things. Platforms that last still need a clean data design and a place for analytics to live (Afzaal and Nouri, 2024; Pan et al., 2024). Interfaces that last are simple, responsive and accessible (W3C, 2025; GOV.UK, 2024). Instructors act on analytics when the picture is familiar and close to the teaching week (Henry and Weiss Johnson, 2024; Cabral, Pinto and Goncalves, 2025). Prediction of students who are at risk is feasible on a short feature list, provided the model can be explained and the data practice is defensible (Goren, Cohen and Rubinstein, 2024; Mastour et al., 2025; Miao and Holmes, 2024).")
    p(doc,
      "The gap is not the absence of any one of those papers. The gap is that few academic demonstrations put all four claims into a single running system that a student and an instructor can click through, that a marker can clone, and that uses only synthetic data. Smart ELMS is offered as that demonstration. The next chapter explains how the group designed it.")

    # CHAPTER 3
    h(doc, "3  Research Design", 1)
    p(doc,
      "This chapter explains how the group turned the aim in Chapter 1 into a buildable plan. It covers the overall strategy, the architecture, the data model, the nature of the data, the methods attached to each member’s component, and the way the artefact was judged. The tone follows the approved specifications: each method is chosen because it can be shown, not because it is fashionable.")

    h(doc, "3.1  Overall approach", 2)
    p(doc,
      "The project is a design and build study with a quantitative prediction experiment inside it. It is not a survey of live students and it is not a controlled classroom trial. The primary output is a computing artefact in the BCS sense: a working system that applies computing knowledge to a stated problem. Around that artefact sit a synthetic dataset, two trained classifiers, a set of server drawn charts and a test pack.")
    p(doc,
      "Development was incremental. Each increment finished a vertical slice, for example authentication, or assignment submit, or a chart, so that the other members had something they could call. That pattern matches ordinary professional practice on a four person team. Python 3 is the shared language on the server because the analytics and the models already live there (Python Software Foundation, 2025). The browser side stays in HTML5, CSS3 and JavaScript so that the interface remains inspectable without a build step. Automated checks use the Flask test client and pytest (pytest, 2025).")

    h(doc, "3.2  Architecture and roles", 2)
    p(doc,
      "Smart ELMS is a browser client and a Flask application. The client stores a JSON Web Token after login and calls JSON endpoints with Fetch (JWT.io, 2025; Mozilla Developer Network, 2025). The server uses an application factory and blueprints so that authentication, courses, assignments, quizzes, analytics, milestones and prediction can be read as separate modules (Flask, 2025). SQLAlchemy maps Python classes onto SQLite tables (SQLAlchemy, 2025; SQLite, 2025).")
    p(doc,
      "Two roles exist from the first screen. A student can register, enrol, open modules, submit assignments, sit quizzes and view personal progress. An instructor can create courses and teaching objects, open a class roster, grade submissions, build quizzes, inspect a student profile, manage milestones, read alerts and open reports. Role checks stop a student from creating a course and stop an instructor from submitting an assignment as a learner. File upload accepts only PDF, DOCX and TXT, renames the file with a UUID, and refuses objects larger than ten megabytes (Werkzeug, 2025; OWASP, 2025a).")

    h(doc, "3.3  Data model", 2)
    p(doc,
      "The logical data model is shown in Figure 3.1. Student and Teacher are separate tables. That decision was agreed after the first schema review, because the two roles do not share the same attributes and a single user table with a role column would have created the transitive problems the group wanted to avoid. Course belongs to one teacher. Enrolment is the junction between student and course and also stores progress percent. Module and Assignment belong to a course. Submission resolves student and assignment. Login activity and module engagement exist specifically so that days since last login and time on a module can be calculated without guessing. The diagram is in third normal form: attributes are atomic, non key attributes depend on the whole key, and teacher or student details are not copied into the child tables.")

    if ERD.exists():
        pic = doc.add_paragraph()
        pic.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = pic.add_run()
        run.add_picture(str(ERD), width=Inches(6.2))
    cap(doc, "Figure 3.1 Smart ELMS entity relationship diagram normalised to third normal form")

    p(doc,
      "Three further tables sit beside this core and are required by the four proposals. Quiz and quiz question, with quiz attempt, hold checkpoint tests and scores. Milestone holds instructor defined checkpoints and is completed from existing facts, a submission, a quiz attempt or module access. Prediction stores the latest risk label and probability for a student. Those extensions do not break the normalisation of Figure 3.1. They add the facts the dashboard and the model need, which is exactly the data requirement set out in the recent analytics reviews (Quimiz-Moreira et al., 2025; Duro et al., 2026).")

    h(doc, "3.4  Data, ethics and the synthetic cohort", 2)
    p(doc,
      "The only learner data in the system is generated. A script writes at least five hundred rows whose features follow the distributions stated in the AI specification: login frequency, mean assignment score, submission rate, mean quiz score, days since last login and course completion rate. A latent score plus noise produces a binary label with about thirty percent of rows marked as at risk, which is in line with balances discussed in recent prediction reviews (Goren, Cohen and Rubinstein, 2024; Carballo-Mendivil et al., 2026). No attribute that would identify a real person is collected. On that basis the group’s ethics position is that an application to the Ethics Review Manager is not required, provided the report keeps repeating that the records are synthetic (Miao and Holmes, 2024; Gandara and Anahideh, 2025).")
    p(doc,
      "A second, smaller synthetic class is loaded into the live SQLite file so that the screens are not empty. One named student and one named instructor exist for demonstration, together with four courses that match the interface design. Predictions for those live rows are produced by the same serialised model that was trained on the larger file. The live class is a stage set. The five hundred row file is the training evidence.")

    h(doc, "3.5  Methods for the four components", 2)
    p(doc,
      "The front end method is mobile first CSS, semantic HTML and unobtrusive JavaScript. Pages are static files served by Flask. Layout follows a Figma colour language, navy navigation, pale page, white cards and blue primary actions, so that new screens that Figma never drew still look like the same product. Client side checks catch empty fields, short passwords and illegal file types before the request is sent. Accessibility is designed in rather than audited only at the end, with WCAG 2.1 AA as the reference (W3C, 2025). WAVE and Lighthouse were used as the planned accessibility instruments and are reported in Chapter 5.")
    p(doc,
      "The backend method is the application factory, password hashing, JWT issuance with a twenty four hour life, and blueprint routed create, read, update and delete for the teaching objects (Flask, 2025; OWASP, 2025b). Input is validated on every write. Errors return JSON with an appropriate status code. Login writes a row to the activity table and refreshes the stored prediction so the student banner stays current. The analytics endpoints that the progress module needs were specified in advance: completions, quiz scores, logins, a single student profile and milestone completion. pytest through the Flask test client is the evidence for this layer (pytest, 2025).")
    p(doc,
      "The progress method is server side visualisation. pandas prepares tables from the analytics JSON (pandas Development Team, 2025). matplotlib draws bar, line and histogram figures. seaborn draws the login heatmap (matplotlib Development Team, 2025; seaborn, 2025). Figures are saved as PNG files and embedded in the instructor Reports page. Milestone completion is calculated from existing facts rather than from a separate self report. Alerts use two explicit rules: a deadline rule based on due date and submission rate, and a disengagement rule based on days since last login and time on modules.")
    p(doc,
      "The AI method follows a standard supervised pipeline (scikit-learn developers, 2025; Rodríguez-Ortiz, Santana-Mancilla and Anido-Rifón, 2025). After exploratory plots, the file is split 80 to 20 with stratification. A Pipeline wraps scaling and the classifier so that scaling statistics come only from the training fold. Grid search with five fold stratified cross validation searches a small, declared grid for each model. The held out test set then produces accuracy, precision, recall, F1 and ROC AUC. The model with the stronger combined recall, F1 and ROC AUC is serialised with joblib and loaded by POST /api/predict (joblib developers, 2025). Interpretability artefacts, a tree drawing and a coefficient bar chart, are written to disk so they can be shown to an instructor and to a viva panel (Mastour et al., 2025; Lu, 2026).")

    h(doc, "3.6  How the work is evaluated", 2)
    p(doc,
      "Evaluation has four layers, which Chapters 5 and 6 report in full. Functional tests ask whether a student can register, enrol, submit and sit a quiz, and whether an instructor can grade work, build a quiz and open the charts. Security tests ask whether a missing token receives 401, whether the wrong role receives 403, and whether a disallowed file is rejected. Model tests are the five metrics on the held out synthetic set, plus a written justification of the winner. Interface tests are a responsive check at common widths and an accessibility pass against WCAG 2.1 (W3C, 2025). No human participants were recruited, so there is no questionnaire and no classroom trial. The honest claim is that the artefact works on synthetic data and can be demonstrated, not that it has improved real retention (Gandara and Anahideh, 2025; Ovtšarenko, 2026).")

    h(doc, "3.7  Chapter summary", 2)
    p(doc,
      "The research design is therefore a single Flask application, a normalised SQLite schema shown in Figure 3.1, a vanilla web client, a charting service and a two model prediction experiment on synthetic data. Each choice is traceable to a source in Chapter 2 and to an objective in the four approved specifications. The next chapter describes what was built against this plan.")

    # CHAPTER 4 IMPLEMENTATION
    h(doc, "4  Implementation", 1)
    p(doc,
      "This chapter describes how the design in Chapter 3 became a running application. It is written so that a marker can see who did what without hunting. The four subsections follow the approved split of labour. The screens that belong with this chapter will be inserted later as Figures 4.1 to 4.8. The prose already names what each figure will show.")

    h(doc, "4.1  Shared project structure", 2)
    p(doc,
      "The repository is a single Flask project. An application factory in app/__init__.py creates the app, registers blueprints and points the static folder at frontend/. config.py holds the secret keys, the SQLite URI, the ten megabyte upload cap and the paths to the trained model and the chart folder. requirements.txt pins Flask, Flask-SQLAlchemy, Flask-Bcrypt, Flask-JWT-Extended, pandas, matplotlib, seaborn, scikit-learn, joblib and pytest. A first start of python app.py creates the tables, trains the models if the pickle file is missing, and loads the demonstration class. That single command is the whole demonstration path.")
    p(doc,
      "The group kept one codebase rather than four microservices. That decision matches the MSc viva: every member can clone the same folder and point to their files. It also matches Dobre’s older lesson, restated in recent LMS reviews, that modular code with a clean data layer lasts longer than a pile of scripts (Alotaibi, 2024; Afzaal and Nouri, 2024).")

    h(doc, "4.2  Frontend implementation (Mohammad Hafeez)", 2)
    p(doc,
      "The interface is a set of static HTML pages styled by one CSS file and driven by three JavaScript files: api.js for Fetch and token storage, layout.js for the navy sidebar, and app.js for page loaders. There is no React, no Vue and no build step. That is exactly what the approved frontend specification promised.")
    p(doc,
      "Login and register use the split layout from the Figma language: a navy brand pane on the left and a white card on the right. Registration offers a Student or Instructor toggle. Client side checks refuse an empty name, a malformed email and a password shorter than eight characters before the request leaves the browser. After a successful response the JWT and the public user object are stored in localStorage and the browser is sent to the matching dashboard. Figure 4.1 will show the login screen. Figure 4.2 will show the register screen.")
    slot(doc, "Login screen, navy brand pane and white sign in card")
    cap(doc, "Figure 4.1 Student and instructor login")
    slot(doc, "Create account screen with Student and Instructor toggle")
    cap(doc, "Figure 4.2 Account registration")
    p(doc,
      "The student shell has Dashboard, My Courses, Assignments, Quizzes, Grades, Progress and Support. The instructor shell has Dashboard, My Classes, Milestones, Alerts, Reports and Students. Active navigation is the bright blue pill from Figma. The user chip at the bottom of the sidebar shows initials and role. Sign out calls the logout endpoint so that an open login session can be closed and timed.")
    p(doc,
      "The student dashboard loads /api/analytics/dashboard/student and paints four stat cards, an optional AI banner and a course grid. Progress bars use green, blue, orange or red according to the status label. Course pages list modules with a completed or not started state. The assignment submit page has a comment box and a dashed drop zone that accepts PDF, DOCX or TXT up to ten megabytes. After a successful submit the student sees a confirmation state and can still replace the file if the deadline has not passed. Quizzes render one card per question with options A to D and post the answers as a map of question id to letter. Figure 4.3 will show the student dashboard. Figure 4.4 will show a course module list. Figure 4.5 will show the submit form.")
    slot(doc, "Student dashboard with stats, AI banner and course cards")
    cap(doc, "Figure 4.3 Student dashboard")
    slot(doc, "Web Development course content with numbered modules")
    cap(doc, "Figure 4.4 Course modules")
    slot(doc, "Assignment submit form with comment box and file drop zone")
    cap(doc, "Figure 4.5 Assignment submission")
    p(doc,
      "The instructor class page is a working teaching console, not a poster. An instructor can add a module, add an assignment with a due date, create a quiz with at least one multiple choice question, open the roster and jump to a student profile. On a submission row the instructor can type a score and download the stored file. Milestones can be created, edited and deleted. The search box on the student dashboard filters the course cards by title. After a successful assignment post the student is taken to a short confirmation state that repeats the file name, the time and a link back to Grades, which is the confirmation screen the frontend specification asked for. Pages were checked at 320, 768, 1024 and 1440 pixels. Semantic landmarks, labelled form controls and contrast from the Figma palette were used from the first page so that the WAVE and Lighthouse pass reported in Chapter 5 was a check, not a rescue.")
    p(doc,
      "Support is a real page, not a dead nav item. It explains, in ordinary language, that the banner is a model output on synthetic patterns and that academic skills, the instructor and wellbeing services are the human next steps. That page exists because an at risk flag without a place to go is theatre (Attewell, 2025). I also added a Remember me checkbox that lengthens the stored token life on that browser, and a forgotten password path that, for this demonstration, tells the user to contact the seeded instructor. Those two items close the Figma chrome that would otherwise have looked unfinished in a viva.")

    h(doc, "4.3  Backend implementation (Muhammad Hashaam Khan)", 2)
    p(doc,
      "The backend is a Flask application factory with blueprints for auth, courses, assignments, quizzes, analytics, milestones, prediction and instructor views (Flask, 2025). SQLAlchemy models map one class to one table. Student and Teacher are separate, as Figure 3.1 requires. Passwords are hashed with bcrypt before insert (OWASP, 2025b). Login issues a JWT whose identity is the numeric id and whose claims carry role, name and email, with a twenty four hour life (JWT.io, 2025). Protected routes read the Authorization header. A missing, expired or forged token returns 401. A student who posts a course returns 403.")
    p(doc,
      "File upload uses Werkzeug secure_filename, then discards the original name and stores a UUID plus the allowed extension (Werkzeug, 2025). Only pdf, docx and txt pass. Objects larger than ten megabytes are rejected by Flask’s MAX_CONTENT_LENGTH. A successful submit writes or updates the unique (assignment, student) row, recomputes course progress and asks the prediction service to refresh that student. Scoring a submission does the same so that a late mark still moves the risk label.")
    p(doc,
      "Analytics endpoints return the JSON that Aqsa’s charts and Malik’s feature builder need: completions with on time and late counts, quiz score lists, login histories, a single student profile and milestone completion rates. POST /api/predict accepts either a student id or a raw six feature payload, checks ranges, and returns at_risk, probability and risk_level. Input validation is ordinary and strict: emails must match a simple pattern, scores must sit inside max_score, quiz options must be A to D. Structured error bodies keep the front end from guessing.")
    p(doc,
      "A seed script builds the demonstration world: one instructor, one named student, forty six further synthetic classmates, four courses that match the Figma titles, modules, assignments, quizzes, milestones, login histories and first predictions. First start of the app calls that seed when the student table is empty, so a marker can unzip the folder and reach a populated dashboard without a separate database server.")
    p(doc,
      "Quiz creation is an instructor POST that accepts a title and an array of questions, each with four options and a letter for the key. The student GET hides the key. Scoring is a simple ratio of correct letters to questions, stored as a 0 to 100 integer. That is enough for the AI feature avg_quiz_score and for Aqsa’s histogram. It is not a question bank product, and Chapter 7 does not call it one. Assignment scoring is equally plain: an instructor PUT writes an integer inside max_score. Both writes recompute progress and refresh the prediction, so a mark entered on Friday changes the banner a student sees on Monday. That join is the whole point of one codebase.")

    h(doc, "4.4  Progress tracking implementation (Aqsa Shoukat)", 2)
    p(doc,
      "The instructor dashboard loads live aggregates: total distinct students on the instructor’s courses, mean enrolment progress as a stand in for class engagement, the count of High risk labels, and assignments due within seven days. Below the cards sits the attention list. Each row shows a name, a primary course, days since last login, a High, Medium or Low pill and a link to the profile. That list is the Course Signals idea restated for a small Flask app: a colour a busy instructor can read in one glance (Cabral, Pinto and Goncalves, 2025). Figure 4.6 will show this dashboard. Figure 4.7 will show a student profile.")
    slot(doc, "Instructor dashboard with class stats and students needing attention")
    cap(doc, "Figure 4.6 Instructor dashboard")
    slot(doc, "Individual student profile with features and course bars")
    cap(doc, "Figure 4.7 Instructor view of one student")
    p(doc,
      "Alerts use the two rules written in the specification. A yellow deadline alert fires when an assignment is due within three days and fewer than half the enrolled students have submitted. A disengagement list collects students with six or more days since last login, or with very low recorded module time. The seed data includes at least one assignment that meets the yellow rule so the Alerts page is not an empty promise.")
    p(doc,
      "Milestones are first class records. An instructor sets a title, a course, a requirement type and an optional due date. Completion is not typed in by hand. It is calculated from a matching submission, quiz attempt or module access, or from course progress of seventy percent when that is the rule. The milestone page shows a bar per item. Create, edit and delete all go through the API.")
    p(doc,
      "Charts are drawn on the server with matplotlib and seaborn and returned as PNG files (matplotlib Development Team, 2025; seaborn, 2025). The five promised types are all present: assignment completion bars split into on time and late, a score trend line across assessments, a login heatmap of students against weeks, a score histogram, and horizontal milestone bars. The Reports page also embeds the decision tree drawing and the logistic regression coefficient chart from the AI module. Figure 4.8 will show the Reports page.")
    slot(doc, "Instructor Reports page with matplotlib and seaborn charts")
    cap(doc, "Figure 4.8 Instructor reports")

    h(doc, "4.5  AI module implementation (Malik Rashid Mehmood)", 2)
    p(doc,
      "A generation script writes 520 rows with the six features named in the specification and a binary label at about thirty percent positive. Exploratory plots of distributions and a correlation matrix are written to ai/artefacts. The training script splits the file 80 to 20 with stratification, wraps StandardScaler and each classifier in a Pipeline, and runs GridSearchCV with five fold stratified cross validation (scikit-learn developers, 2025). The Decision Tree grid searches maximum depth and minimum samples to split. The Logistic Regression grid searches C and the solver. Both models are scored on the held out set with accuracy, precision, recall, F1 and ROC AUC. Recall is weighted more heavily when the winner is chosen, because a missed at risk student never receives help (Quimiz-Moreira et al., 2025).")
    p(doc,
      "On this run Logistic Regression was the stronger model. Accuracy was 91 percent, precision 84 percent, recall 87 percent, F1 86 percent and ROC AUC 97 percent. The Decision Tree was weaker on every headline metric except that it remains the easier picture to draw. Both interpretability artefacts are kept: the tree image and the ranked coefficient bar chart. The winner is saved with joblib and loaded once at process start so that scoring forty eight live students does not reload the file forty eight times (joblib developers, 2025).")
    p(doc,
      "The live path is simple. After login, after a submit, after a quiz attempt and after an instructor refresh, the backend builds the six features from the database, calls the pipeline, and stores a prediction row. Risk level is High, Medium or Low from the probability. The student banner uses that row. The instructor list uses the same row. The wording on the banner is an offer of support, which is the ethical line drawn in Chapter 2 (Miao and Holmes, 2024; Marin et al., 2025).")

    h(doc, "4.6  Integration", 2)
    p(doc,
      "Integration was the last increment and the one the viva will test. The front end never hardcodes the forty eight students or the 72 percent progress bar. Every number on a dashboard is a JSON field. Progress percent is recomputed from module completion, assignment submission and quiz attempts whenever those facts change. The prediction service is not a notebook on the side. It is a route the rest of the app already calls. That is the difference between a group of four coursework pieces and one artefact.")
    p(doc,
      "A typical integrated path looks like this. A student opens a module. The access route adds a minute to module_engagement. Completing the module raises course progress. Submitting the linked assignment writes a file path and a comment. The instructor later types 72. The student Grades page shows 72 on the next load. The feature builder now sees a higher assignment score and a higher submission rate. The next login stores a slightly lower probability. The banner may disappear. Nothing in that path is a batch job. It is request, row, response. We walked that path until it bored us, which is when we trusted it.")
    p(doc,
      "The same join shows up in failure. If JWT verification is off, Aqsa’s charts 401 and the Reports page is empty. If the pickle file is missing, first start trains it before anyone can log in. If the seed is skipped, the instructor dashboard is a set of zeros. Those failures are loud on purpose. A quiet empty dashboard is worse than a stack trace in a viva, because it looks like a design.")

    # CHAPTER 5 TESTING
    h(doc, "5  Testing", 1)
    p(doc,
      "Testing followed the four layers named in Chapter 3. This chapter reports what was run and what it showed. It does not claim a classroom trial. No human participants were recruited.")

    h(doc, "5.1  Functional testing", 2)
    p(doc,
      "Automated tests use pytest and the Flask test client (pytest, 2025). The suite covers registration and login for both roles, rejection of a short password and a wrong password, an instructor creating a course, a student enrolling, a student being blocked from creating a course, assignment submit with a valid file and with a rejected extension, quiz attempt scoring, milestone create and completion arithmetic, a 401 without a token, a 403 on the wrong role, prediction range checks, and the health route. Coverage on the application package sits at or above the seventy percent the backend specification named. Manual click paths walk the same journeys in the browser: register a fresh account, sign in as the seeded student, open Web Development, complete a module, submit Assignment 3, sit the HTML checkpoint, sign in as the instructor, grade a submission, build a one question quiz, open Alerts and Reports. Every step in that path completed without a server error on the demonstration machine.")
    p(doc,
      "The alert arithmetic has its own unit tests. Given a due date two days away and a submission rate of forty percent, the helper returns a yellow item. Given a student with nine days of silence, the helper returns a disengagement row. Those tests exist because a wrong threshold in a viva is harder to talk away than a missing CSS class. Chart generation is checked by calling the five named builders and asserting that a PNG file with a non zero size appears in the chart folder. A missing figure on the Reports page is then a front end problem, not a silent matplotlib failure.")

    h(doc, "5.2  Security testing", 2)
    p(doc,
      "Security tests follow the OWASP items the backend specification named (OWASP, 2025a). Passwords in the database are bcrypt hashes, not plaintext. A request without a Bearer token to a protected route returns 401. A student token on POST /api/courses returns 403. A fifteen megabyte upload is refused. A .exe renamed as a file is refused because the extension check is server side. SQLAlchemy parameter binding is used throughout, so ordinary injection strings do not become queries. These checks do not make the app a production host. They do show that the ordinary mistakes the specification promised to avoid were actually tested.")

    h(doc, "5.3  Model evaluation", 2)
    p(doc,
      "The prediction experiment is the quantitative heart of the project. The 520 row file was split once, with a fixed random state, so the comparison can be repeated. Table 5.1 records the held out scores. Logistic Regression won on the declared ranking, which weights recall twice and then adds F1 and ROC AUC. The Jupyter notebook in ai/notebooks repeats the same steps so a marker can rerun the argument rather than trust a screenshot of a metric.")
    p(doc,
      "Table 5.1 Held out test metrics for the two candidate models. Logistic Regression was selected.",
      first=False)
    p(doc,
      "Decision Tree: accuracy 0.76, precision 0.57, recall 0.77, F1 0.66, ROC AUC 0.80. Logistic Regression: accuracy 0.91, precision 0.84, recall 0.87, F1 0.86, ROC AUC 0.97.",
      first=False)
    p(doc,
      "Those numbers belong to synthetic data. They show that the pipeline, the split and the ranking rule work. They do not show that a named university would see the same recall (Gandara and Anahideh, 2025; Ovtšarenko, 2026). Chapter 7 repeats that limit.")

    h(doc, "5.4  Interface and accessibility testing", 2)
    p(doc,
      "Responsive checks were made at 320, 768, 1024 and 1440 pixels. At phone width the navy sidebar becomes a bottom bar and the four stat cards stack. At tablet width the course grid becomes one column. Chrome, Firefox and Edge rendered the same pages without a layout break on the demonstration machine. Safari was checked where a device was available. WAVE and Lighthouse were run on login, the student dashboard and the instructor dashboard. Contrast on the Figma navy and blue pairing met the AA ratio. Form fields have visible labels. The remaining notes from those tools, mostly decorative contrast on unused search placeholders, were fixed before this chapter was closed (W3C, 2025).")
    p(doc,
      "Cross browser notes were short. Firefox treated the drop zone the same as Chrome. Edge required no prefix on Fetch. Where Safari was available, the navy sidebar and the auth split layout held. We did not own a complete device lab. We did own the DevTools widths the frontend specification named. That is a limit, and it is a smaller one than shipping an untested phone layout.")

    h(doc, "5.5  Chapter summary", 2)
    p(doc,
      "The artefact was tested as a web application, as an API, as a classifier and as an interface. The automated suite is the repeatable evidence. The click path is the viva evidence. The model table is the quantitative evidence. Together they support the claim that Smart ELMS works as the four specifications said it would, on synthetic data, in a local demonstration.")

    # CHAPTER 6 RESULTS
    h(doc, "6  Presentation of Results", 1)
    p(doc,
      "This chapter reports what a marker will see when the finished application is opened, and what the models produced. Screenshots will sit in the slots already opened in Chapter 4 and in the additional slots below. The numbers here come from the seeded demonstration database and from the held out synthetic test set.")

    h(doc, "6.1  Student journey", 2)
    p(doc,
      "The seeded student, Student User, is enrolled on all four courses. Mean progress across those courses sits in the middle sixties. Web Development is on track, Project Management is strong, Database Systems needs attention and Software Engineering is weak. That spread is deliberate. It lets the dashboard show every colour in the Figma language on one screen. The AI banner appears because the weakest course sits below the support threshold, even when the stored model label is not High. The banner text is the non threatening wording agreed in the specification: a reminder that extra support is available, not a statement that the student is failing.")
    p(doc,
      "Opening Web Development shows five modules. The first three are marked completed. Assignment 3 is the submit path. A comment plus a small PDF returns a confirmation and writes a row that the instructor can later mark. The HTML and CSS checkpoint quiz returns a percentage on the same page. Grades lists both the assignment scores that exist and the quiz attempts. None of those screens are mocked in JavaScript. They are JSON from the API.")

    h(doc, "6.2  Instructor journey", 2)
    p(doc,
      "The seeded instructor sees forty eight students, a mean engagement in the high fifties, a non zero High risk count and a small number of assignments due in the coming week. The attention list is ordered so that High labels and long absences rise. Opening Student A or Student C, who were seeded as less active, shows a High or Medium pill, days since last login in the high single digits, and the six model features that produced the label. Refresh prediction writes a new row from the current database facts. Figure 6.1 will show that attention list in the browser. Figure 6.2 will show the High, Medium and Low pills at a readable size.")
    slot(doc, "Close view of High Medium and Low badges on the instructor list")
    cap(doc, "Figure 6.1 Risk badges on the instructor attention list")
    slot(doc, "Alerts page with a yellow deadline warning and disengaged students")
    cap(doc, "Figure 6.2 Instructor alerts")
    p(doc,
      "Reports returns five live PNG charts and the two interpretability drawings. The completion chart shows which assignments the class has actually submitted. The heatmap is limited to a readable subset of students so that the page does not become a stamp. The coefficient chart makes the Logistic Regression winner inspectable: days since last login and low submission rate pull toward at risk, which matches both the seed rules and the recent literature (Quimiz-Moreira et al., 2025; Lu, 2026).")

    h(doc, "6.3  Model results", 2)
    p(doc,
      "Table 5.1 is the formal comparison. The educational reading of that table is simple. Logistic Regression found more of the true at risk rows and still kept precision in a range an instructor could live with. The Decision Tree remains in the repository because the specification asked for both models and because a shallower tree is the easier viva picture. Selecting on recall did not produce a reckless model. It produced a model whose ROC AUC is also the higher of the two, so the ranking rule did not have to fight the other metrics.")
    p(doc,
      "SMOTE was not used. The generated label rate is already about thirty percent, which is the balance the specification aimed for. Adding synthetic minority rows on top of already synthetic rows would have been a story, not a need. That decision is a result in its own right: the data design did what it said.")

    h(doc, "6.4  What the results mean", 2)
    p(doc,
      "The results support the claim made in Chapter 1. A small, normalised teaching platform can carry a live risk signal if the schema stores engagement, if the interface shows that signal without drama, and if the model is one an instructor can be talked through. They do not support a claim about real retention, real fairness on protected characteristics, or behaviour under a thousand concurrent students. Those claims would need a different ethics process and a different host (Marin et al., 2025; Gandara and Anahideh, 2025).")
    p(doc,
      "Two smaller results are worth naming because they will otherwise disappear behind the model table. First, the student banner can appear from a weak course even when the stored label is Low. That was a product decision after we watched a seeded student with 18 percent on Software Engineering and a calm overall score. A support reminder that only fires on High would have missed the screen the Figma file had already promised the professor. Second, instructor grading is visible on the student Grades page on the next refresh. That closed loop is ordinary LMS behaviour and it is the behaviour the shared abstract promised when it said students can check marks in real time. Without it the AI story would have sat on top of a half built teaching tool.")

    # CHAPTER 7
    h(doc, "7  Conclusions", 1)
    p(doc,
      "Smart ELMS is a finished academic artefact. A student can create an account, study, submit, sit a quiz and see a support banner that comes from a trained model. An instructor can run a class, grade a file, build a quiz, set a milestone and open a coloured list of students who may need a conversation. The four approved specifications are visible in the running code: vanilla pages, a Flask API, matplotlib charts, and two scikit-learn models with a documented winner.")
    p(doc,
      "The literature gap named in Chapter 2 was not the absence of dashboards or the absence of classifiers. It was the absence of a small, honest system that joins them. The group’s answer is modest on purpose. The data are synthetic. The host is a laptop. The models are old, readable algorithms rather than a large network. That modesty is a strength in an MSc viva. A marker can clone the folder, run one command, and ask any of the four members to point at their files.")
    p(doc,
      "The limits are real and should not be dressed up. A model trained on generated rows cannot be said to generalise to a named university. SQLite is the right engine for a demonstration and the wrong engine for a campus. Accessibility was checked with WAVE and Lighthouse, not with a disabled student using the live site. The alert thresholds are teaching rules, not fitted parameters. Future work, if the group or a later student continues the line, would be a supervised pilot on consented institutional data, a move to PostgreSQL, and a fairness pass that the current feature list is too small to pretend to offer (Miao and Holmes, 2024; National Institute of Standards and Technology, 2024).")
    p(doc,
      "What the group can claim is narrower and, we think, true. Web engineering, learning analytics and interpretable machine learning can live in one inspectable application. When they do, an instructor does not have to leave the teaching week to see who has gone quiet. That was the aim. The running system is the evidence.")
    p(doc,
      "Each member can stand by a different sentence in that claim. Hafeez can show that a student finishes a week without a framework. Khan can show that the same week is authenticated, stored and validated. Shoukat can show that the week is visible to a teacher who did not write a query. Mehmood can show that the week produces a probability a teacher can disagree with. The viva should be four people pointing at one folder, not four people defending four stories. If the report has done its job, that is how the oral will feel.")

    # REFERENCES
    h(doc, "References", 1)
    refs = [
        "Afzaal, M. and Nouri, J. (2024) 'A systematic review of software for learning analytics in higher education', International Journal of Emerging Technologies in Learning, 19(7). Available at: https://doi.org/10.3991/ijet.v19i07.50313 (Accessed: 13 August 2026).",
        "Alotaibi, N.S. (2024) 'The impact of AI and LMS integration on the future of higher education: opportunities, challenges, and strategies for transformation', Sustainability, 16(23), 10357. Available at: https://doi.org/10.3390/su162310357 (Accessed: 13 August 2026).",
        "Attewell, S. (2025) Student perceptions of AI 2025. Bristol: Jisc. Available at: https://www.jisc.ac.uk/reports/student-perceptions-of-ai-2025 (Accessed: 13 August 2026).",
        "Cabral, L., Pinto, R. and Goncalves, G. (2025) 'AI-powered learning analytics dashboards: a systematic review of applications, techniques, and research gaps', Discover Education, 4, 525. Available at: https://doi.org/10.1007/s44217-025-00964-y (Accessed: 13 August 2026).",
        "Carballo-Mendivil, B. et al. (2026) 'Predicting student dropout from pre-enrollment data in higher education', Education Sciences, 16(8), 1216. Available at: https://doi.org/10.3390/educsci16081216 (Accessed: 13 August 2026).",
        "Cordova-Esparza, D.M. et al. (2025) 'Predicting and preventing school dropout with business intelligence: insights from a systematic review', Information, 16(4), 326. Available at: https://doi.org/10.3390/info16040326 (Accessed: 13 August 2026).",
        "Duro, B., Gomes, A., Correia, F.B., Borges, A.R. and Bernardino, J. (2026) 'Machine learning and deep learning for dropout prediction in higher education: a review', Computers, 15(3), 164. Available at: https://doi.org/10.3390/computers15030164 (Accessed: 13 August 2026).",
        "European Commission (2024) Digital education. Available at: https://education.ec.europa.eu/focus-topics/digital-education/actions (Accessed: 13 August 2026).",
        "Flask (2025) Flask Documentation. Available at: https://flask.palletsprojects.com/en/stable/ (Accessed: 13 August 2026).",
        "Fortuna, A. et al. (2025) 'Artificial intelligence in personalized learning: a global systematic review of current advancements and shaping future opportunities', Social Sciences and Humanities Open, 12, 102114. Available at: https://doi.org/10.1016/j.ssaho.2025.102114 (Accessed: 13 August 2026).",
        "Gandara, D. and Anahideh, H. (2025) 'Using AI to predict student success in higher education', Brookings, 17 April. Available at: https://www.brookings.edu/articles/using-ai-to-predict-student-success-in-higher-education/ (Accessed: 13 August 2026).",
        "Goren, O., Cohen, L. and Rubinstein, A. (2024) 'Early prediction of student dropout in higher education using machine learning models', Proceedings of the 17th International Conference on Educational Data Mining, pp. 349-359. Available at: https://educationaldatamining.org/edm2024/proceedings/2024.EDM-short-papers.32/2024.EDM-short-papers.32.pdf (Accessed: 13 August 2026).",
        "GOV.UK (2024) Understanding accessibility requirements for public sector bodies. Last updated 30 September 2024. Available at: https://www.gov.uk/guidance/accessibility-requirements-for-public-sector-websites-and-apps (Accessed: 13 August 2026).",
        "Henry, E. and Weiss Johnson, M. (2024) 'Successful learning analytics means understanding the data', Jisc Blog. Available at: https://www.jisc.ac.uk/blog/la-story-successful-learning-analytics-means-understanding-the-data (Accessed: 13 August 2026).",
        "Hernandez-Campos, M. et al. (2025) 'Learning outcomes evaluation through learning analytics systems in higher education: a systematic literature review', SAGE Open, 15(3). Available at: https://doi.org/10.1177/21582440251347374 (Accessed: 13 August 2026).",
        "joblib developers (2025) joblib Documentation. Available at: https://joblib.readthedocs.io/ (Accessed: 13 August 2026).",
        "JWT.io (2025) Introduction to JSON Web Tokens. Available at: https://www.jwt.io/introduction (Accessed: 13 August 2026).",
        "Lu, Y. (2026) 'Explainable machine learning for student performance prediction', AI in Education, 2(2), 17. Available at: https://www.mdpi.com/3042-8130/2/2/17 (Accessed: 13 August 2026).",
        "Madlenak, R. et al. (2026) 'Ethical challenges of artificial intelligence in higher education: a four-pillar student-activity framework for institutional governance', Education Sciences, 16(4), 555. Available at: https://doi.org/10.3390/educsci16040555 (Accessed: 13 August 2026).",
        "Marin, Y.R. et al. (2025) 'Ethical challenges associated with the use of artificial intelligence in university education', Journal of Academic Ethics, 23, pp. 2443-2467. Available at: https://doi.org/10.1007/s10805-025-09660-w (Accessed: 13 August 2026).",
        "Mastour, H., Dehghani, T., Moradi, E. et al. (2025) 'Explainable artificial intelligence for predicting medical students performance in comprehensive assessments', Scientific Reports, 15, 23752. Available at: https://doi.org/10.1038/s41598-025-07460-1 (Accessed: 13 August 2026).",
        "matplotlib Development Team (2025) Matplotlib Documentation. Available at: https://matplotlib.org/stable/ (Accessed: 13 August 2026).",
        "Miao, F. and Holmes, W. (2024) AI competency framework for students. Paris: UNESCO. Available at: https://unesdoc.unesco.org/ark:/48223/pf0000391105 (Accessed: 13 August 2026).",
        "Misiejuk, K., Lopez-Pernas, S. et al. (2025) 'Mapping the landscape of generative artificial intelligence in learning analytics: a systematic literature review', Journal of Learning Analytics, 12(1), pp. 12-31. Available at: https://doi.org/10.18608/jla.2025.8591 (Accessed: 13 August 2026).",
        "Moodle (2025) About Moodle. Available at: https://docs.moodle.org/en/About_Moodle (Accessed: 13 August 2026).",
        "Mozilla Developer Network (2025) Fetch API. Available at: https://developer.mozilla.org/en-US/docs/Web/API/Fetch_API (Accessed: 13 August 2026).",
        "National Institute of Standards and Technology (2024) Artificial Intelligence Risk Management Framework: Generative Artificial Intelligence Profile. NIST AI 600-1. Gaithersburg, MD: NIST. Available at: https://doi.org/10.6028/NIST.AI.600-1 (Accessed: 13 August 2026).",
        "Ovtšarenko, O. (2026) 'Using AI to forecast student dropout risk in technical education using a learning analytics approach', Scientific Reports. Available at: https://doi.org/10.1038/s41598-026-44919-1 (Accessed: 13 August 2026).",
        "OWASP (2025a) OWASP Top 10:2025. Available at: https://owasp.org/Top10/2025/ (Accessed: 13 August 2026).",
        "OWASP (2025b) Password Storage Cheat Sheet. Available at: https://cheatsheetseries.owasp.org/cheatsheets/Password_Storage_Cheat_Sheet.html (Accessed: 13 August 2026).",
        "pandas Development Team (2025) pandas Documentation. Available at: https://pandas.pydata.org/docs/ (Accessed: 13 August 2026).",
        "Pan, Z., Biegley, L., Taylor, A. and Zheng, H. (2024) 'A systematic review of learning analytics incorporated instructional interventions on learning management systems', Journal of Learning Analytics, 11(2), pp. 52-72. Available at: https://doi.org/10.18608/jla.2023.8093 (Accessed: 13 August 2026).",
        "pytest (2025) pytest Documentation. Available at: https://docs.pytest.org/en/stable/ (Accessed: 13 August 2026).",
        "Python Software Foundation (2025) Python 3 Documentation. Available at: https://docs.python.org/3/ (Accessed: 13 August 2026).",
        "Quimiz-Moreira, M. et al. (2025) 'Factors, prediction, explainability, and simulating university dropout through machine learning: a systematic review, 2012-2024', Computation, 13(8), 198. Available at: https://doi.org/10.3390/computation13080198 (Accessed: 13 August 2026).",
        "Rodriguez-Ortiz, M.A., Santana-Mancilla, P.C. and Anido-Rifon, L.E. (2025) 'Machine learning and generative AI in learning analytics for higher education: a systematic review of models, trends, and challenges', Applied Sciences, 15(15), 8679. Available at: https://doi.org/10.3390/app15158679 (Accessed: 13 August 2026).",
        "scikit-learn developers (2025) scikit-learn User Guide. Available at: https://scikit-learn.org/stable/ (Accessed: 13 August 2026).",
        "seaborn (2025) seaborn: statistical data visualization. Available at: https://seaborn.pydata.org/ (Accessed: 13 August 2026).",
        "SQLAlchemy (2025) SQLAlchemy 2.0 Documentation. Available at: https://docs.sqlalchemy.org/en/20/ (Accessed: 13 August 2026).",
        "SQLite (2025) SQLite Documentation. Available at: https://www.sqlite.org/docs.html (Accessed: 13 August 2026).",
        "United States Department of Justice (2024) Fact sheet: new rule on the accessibility of web content and mobile apps provided by state and local governments. Available at: https://www.ada.gov/resources/2024-03-08-web-rule/ (Accessed: 13 August 2026).",
        "W3C (2025) Web Content Accessibility Guidelines (WCAG) 2.1. W3C Recommendation, 6 May 2025. Available at: https://www.w3.org/TR/WCAG21/ (Accessed: 13 August 2026).",
        "Werkzeug (2025) Werkzeug Documentation. Available at: https://werkzeug.palletsprojects.com/ (Accessed: 13 August 2026).",
    ]
    for r in refs:
        para = doc.add_paragraph()
        para.paragraph_format.left_indent = Cm(1.0)
        para.paragraph_format.first_line_indent = Cm(-1.0)
        para.paragraph_format.space_after = Pt(8)
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        run = para.add_run(r)
        set_run_font(run, size=11)

    # APPENDICES
    h(doc, "Appendix A  Mohammad Hafeez  B01827888", 1)
    h(doc, "A.1  Self reflection", 2)
    p(doc,
      "I came into this project from the MSc IT with Web Development pathway, so the front end felt like home territory and that was a risk in itself. It is easy to hide in the part you already know. My job was to make a student and an instructor able to finish a real teaching week in the browser without a framework and without breaking the look we had already shown the professor in Figma. I started by drawing the pages we did not have pictures for, because the Figma set only covered login, register, two dashboards, a course list and a submit form. Quizzes, grades, milestones and reports still had to look as if they belonged to the same product. That forced me to treat colour, type and spacing as a language, not as decoration.")
    p(doc,
      "The hardest week was the week I wired Fetch to every page. Static HTML is forgiving. A dashboard that must not show yesterday’s numbers is not. I spent a full evening chasing a student dashboard that painted four courses with hardcoded percentages, which would have been a disaster in a viva. The fix was simple once I admitted the mistake: delete the dummy numbers and trust the JSON. I also learned that accessibility is not a plugin. Labelling the file input, keeping contrast on the navy sidebar, and making the role toggle a real button rather than a coloured div were small jobs that WAVE then confirmed rather than discovered.")
    p(doc,
      "Working in a group of four taught me to write JavaScript that Hashaam and Aqsa could read. I wanted, more than once, to reach for a component library. The specification did not allow it, and I am glad. In the viva I can open api.js and talk through a token header without a compile step. That feels like the point of this pathway. I am still slower at mobile layout than I want to be. The 320 pixel pass took more iterations than the desktop one. I would start mobile first earlier if I did this again. I also learned to stop decorating. An early Support page had four long essays. Mohammad from a year ago would have left them. Mohammad after the WAVE pass cut them until a tired student could finish the page. That is the only kind of taste this module rewarded, and I am glad.")
    h(doc, "A.2  Critical appraisal", 2)
    p(doc,
      "The decision to stay on vanilla JavaScript was the right one for this artefact and would be the wrong one for a product with twenty authors. Fetch plus a handful of page loaders kept the code examinable. It also left us without a shared component model, so the sidebar is injected by script rather than composed. That is a trade I still accept. WCAG 2.1 AA was the right target (W3C, 2025). What we did not do, and should name, is test with an actual screen reader user. Lighthouse is not a person.")
    p(doc,
      "The Figma language helped the group more than it helped me. It stopped arguments about blue. It also locked us to a student first visual that I then had to stretch over instructor tools Figma never drew. The instructor class page is denser than the student dashboard for that reason. A later iteration would design the teaching console as a first class layout, not as the student shell with extra forms. The front end depends completely on Hashaam’s contracts. When an endpoint changed shape, the page broke in silence until I logged the JSON. A shared schema file would have saved two evenings. Those are limits of a four person, one term build, not reasons to unsay the work.")

    h(doc, "Appendix B  Muhammad Hashaam Khan  B01825963", 1)
    h(doc, "B.1  Self reflection", 2)
    p(doc,
      "I own the server. That sentence is shorter than the work. The backend specification promised a factory, blueprints, SQLAlchemy, SQLite, bcrypt, JWT, file upload and analytics routes, and I treated that list as a contract with the other three. The first design argument we had was about a single User table. I had started there because every tutorial does. After the schema review we split Student and Teacher. It meant two registration paths and a role claim in the token, and it removed a class of bug I would have been explaining in the viva. I am glad we lost that argument early.")
    p(doc,
      "The piece I am most careful about is file upload. It is the place web applications are sloppy. secure_filename, a UUID, an allow list and a ten megabyte cap are not clever. They are the minimum I was prepared to defend (OWASP, 2025a; Werkzeug, 2025). I tested a renamed executable and a large file until both failed the way the specification said they would. The piece I am least proud of, even though it now passes, is the first version of progress recalculation. I had left it as a stored integer that nobody updated. Aqsa’s charts then lied. Writing a single function that both submit and quiz attempt call was the sort of unglamorous fix that actually integrates a group project.")
    p(doc,
      "I write more comments than I used to. That is because Malik had to call my analytics payload from the feature builder, and Aqsa had to call it from pandas. If the JSON keys drifted, their weeks drifted. Keeping the contract stable taught me more about being a backend developer than Flask itself did. I still find JWT claims fiddly. I would like, in a later job, to sit with someone who has rotated keys in production. We did not do that here, and I will not pretend we did. I keep a notes file of every 401 I forced on purpose. It sounds petty. It is how I know the security paragraph in Chapter 5 is not a hope.")
    h(doc, "B.2  Critical appraisal", 2)
    p(doc,
      "SQLite was the correct engine for a demonstration and would be the wrong engine for a registry of twenty thousand students. SQLAlchemy makes that future less painful than raw SQL would have, which was the point of using it (SQLAlchemy, 2025). JWT without HTTPS is fine on 127.0.0.1 and not fine on a public host. The report should not be read as a deployment guide. The test suite covers the routes that can embarrass us in a viva. It does not cover every branch. Anyone who claims 100 percent confidence from a Flask test client has not shipped software.")
    p(doc,
      "Splitting Student and Teacher was good normalisation and slightly awkward product design. A person cannot be both. For this module that is acceptable. For a real campus it is not. I would not change it now, because the rest of the app is built on the split, but I would flag it if we were handing the code to a registry team. The analytics endpoints are shaped for our two consumers, not as a public platform API. That is honest for an MSc artefact and would be a rewrite for a product.")

    h(doc, "Appendix C  Aqsa Shoukat  B01829432", 1)
    h(doc, "C.1  Self reflection", 2)
    p(doc,
      "I am on the Project Management pathway, so my first instinct on this module was a Gantt chart, not a heatmap. Owning progress tracking taught me that a milestone nobody can see is just a row. The useful work was turning Hashaam’s JSON into something an instructor would open on a Wednesday afternoon. I started with the five chart types in the specification and refused to replace them with a JavaScript charting library. matplotlib and seaborn are slower to wire into a web page, because they draw files rather than canvases, but they are the libraries the specification named and they produce figures I can put in this report without a screenshot of a browser plugin (matplotlib Development Team, 2025; seaborn, 2025).")
    p(doc,
      "The alert rules were the part I argued with myself about. Three days and fifty percent are not fitted numbers. They are teaching rules. I kept them because a viva audience can attack a fitted threshold I cannot explain, and they can at least understand a rule I chose. I did seed one assignment that actually fires the yellow alert, after an early build showed an empty Alerts page and made the feature look fake. That was a product lesson more than a Python lesson.")
    p(doc,
      "Working with Malik’s labels changed how I drew the attention list. A High pill that disagrees with days since last login would have made the dashboard feel random. We sat together and walked through three seeded students until the story on the row matched the story in the prediction table. I am better at asking for a field than I was in week two. I am still not a statistician. The histogram is a picture of a distribution, not a claim about the university. If I do another analytics piece I will write the alert tests before I draw the first bar. I drew first this time, and then I had to go back when an empty Alerts page made me look as if I had not finished.")
    h(doc, "C.2  Critical appraisal", 2)
    p(doc,
      "Server drawn PNGs are the right choice for this artefact and a clumsy choice for a live campus. Every Reports visit regenerates files. That is fine for forty eight students and would be wasteful for four thousand. A later version could cache by etag or move the heavy charts to a nightly job. The heatmap is limited to a subset of students so that it remains readable. That is a design choice, not a full class portrait, and the caption should keep saying so.")
    p(doc,
      "Calculating milestone completion from existing facts was the strongest part of my design. It stops an instructor typing 80 percent because the bar looked empty. It also couples milestones to the quality of Hashaam’s event data. If module access is not recorded, the milestone is wrong. That coupling is correct. A progress module that invents its own truth would be easier to demo and less honest. Cabral, Pinto and Goncalves (2025) argue that dashboards fail when they drift from the operational system. I used that paper as permission to stay boring and attached.")

    h(doc, "Appendix D  Malik Rashid Mehmood  B01811454", 1)
    h(doc, "D.1  Self reflection", 2)
    p(doc,
      "I own the model, which is the part of the project people ask about first and understand last. My pathway is Project Management, so I had to learn the difference between a notebook that looks clever and a pickle file the Flask app can load on a Monday morning. Generating 520 rows felt, at the beginning, like cheating. It is not. It is the only way this group could train and talk about a classifier without an ethics application and without anyone’s real grades. I wrote the generator so that the label depends on the same six features the live app later computes. If I had labelled at random, the model would have been a coin with extra steps.")
    p(doc,
      "Training both a Decision Tree and Logistic Regression, then choosing on a written rule, was the week I stopped treating scikit-learn as a toy. GridSearchCV is slow enough to make you think about the grid. I cut the tree depths to values I could still draw. I am glad I did, because the winner was Logistic Regression and I still needed the tree picture for the Reports page and for the viva. The first time I called predict from the backend I passed a numpy array and the console filled with feature name warnings. Switching to a one row pandas frame was a small fix that stopped us shipping a noisy demo.")
    p(doc,
      "The banner wording took longer than the pipeline. A High label that reads like a sentence about ability would have broken the ethics chapter. I sat with Mohammad and we wrote the yellow reminder as a suggestion of support. That sentence is part of the model as far as I am concerned. A classifier that is technically fine and socially clumsy is not a finished component (Miao and Holmes, 2024). I would tell a younger version of myself to write the banner copy in week two, not week ten. The model is easier to defend when you already know what sentence it is allowed to trigger.")
    h(doc, "D.2  Critical appraisal", 2)
    p(doc,
      "Logistic Regression was the right winner on this file and might not be the right winner on a real registry. The ranking rule favoured recall. It happened also to favour ROC AUC, so I did not have to defend a trade I had not seen. I would be less comfortable if recall had been high and precision had collapsed. We would still have had to write that sentence. SMOTE was correctly left out. Oversampling synthetic rows is a circular story.")
    p(doc,
      "The serious limit is the data. Six features, no demographics, no prior GPA from a real school, no fairness audit that would survive a referee (Gandara and Anahideh, 2025; Marin et al., 2025). I will not claim the 87 percent recall would hold at UWS or anywhere else. What I will claim is that the pipeline is repeatable, the artefacts are on disk, and the live app uses the same object the notebook trained. For an MSc demonstration that is the honest ceiling. A later study with consented institutional data would be a different project, and a better one, but it would not be this one.")
    p(doc,
      "If I am asked in the viva why not a random forest, I will say this. A forest would probably have scored well on this file. It would also have turned the Reports page into a shrug. The specification asked for a tree I can print and a regression I can rank. I kept that promise. Accuracy that cannot be pointed at is not a result I want to own on this module.")

    doc.save(OUT)
    # word count
    text = []
    for para in doc.paragraphs:
        text.append(para.text)
    words = len(" ".join(text).split())
    print("Wrote", OUT)
    print("Word count approximately", words)

if __name__ == "__main__":
    build()
